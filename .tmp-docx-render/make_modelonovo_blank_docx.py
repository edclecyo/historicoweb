from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "modelonovo.docx"
OUTPUT = ROOT / ".tmp-docx-render" / "modelonovo-blank-docx-latest.docx"


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def clear_cell_at(table, row: int, col: int) -> None:
    try:
        clear_cell(table.rows[row - 1].cells[col - 1])
    except IndexError:
        return


def clear_runs_text(document: Document, samples: tuple[str, ...]) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for sample in samples:
                if sample in run.text:
                    run.text = run.text.replace(sample, "")


def clear_paragraph_containing(document: Document, needle: str) -> None:
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            for run in paragraph.runs:
                run.text = ""


document = Document(INPUT)

student_table = document.tables[1]
for col in range(1, 5):
    clear_cell_at(student_table, 2, col)

main_table = document.tables[2]

for row in range(2, 4):
    for col in range(3, 12):
        clear_cell_at(main_table, row, col)

for row in range(4, 15):
    for col in range(4, 13):
        clear_cell_at(main_table, row, col)

for row in range(16, 27):
    for col in range(2, 11):
        clear_cell_at(main_table, row, col)

for row in range(27, 33):
    for col in range(1, 11):
        clear_cell_at(main_table, row, col)

for col in range(2, 11):
    clear_cell_at(main_table, 33, col)
    clear_cell_at(main_table, 34, col)

for row in range(35, 38):
    for col in range(3, 12):
        clear_cell_at(main_table, row, col)

for row in range(40, 49):
    for col in range(3, 7):
        clear_cell_at(main_table, row, col)

clear_cell_at(document.tables[3], 2, 1)

clear_runs_text(
    document,
    (
        "TESTE DE COMO DEVE SER",
        "TESTE PAI",
        "TESTE MAE",
        "TESTE DO NOME COMO DEVE SER NO CERTIFICADO",
        "9� ANO",
        "9° ANO",
        "9º ANO",
        "2026",
        "ENSINO MEDIO",
    ),
)
clear_paragraph_containing(document, "Brejo Santo- Cear")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(OUTPUT)
